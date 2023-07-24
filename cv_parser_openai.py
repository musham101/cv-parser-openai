import base64
import requests
import fitz, json
from docx import Document
from globals import API_KEY

api_key = API_KEY

q_one_liner = "get me the name, email, phone, city, educations, experiences and skills in json format from the given text"

class CV_PARSER:

    # Convert Word file to base64.
    def word_to_base64(doc_path):
        doc = Document(doc_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        base64_encoded = base64.b64encode(text.encode('utf-8')).decode('utf-8')
        return base64_encoded
    
    # Convert base64 to word file.
    def base64_to_word(base64_string):
        decoded_bytes = base64.b64decode(base64_string)
        text = decoded_bytes.decode('utf-8')
        doc = Document()
        doc.add_paragraph(text)
        doc.save('resume\last_parsed_cv.docx')
        print(f"Word document has been successfully created at 'last_parsed_cv.docx'.")
    

    # Convert pdf file to base64.
    def pdf_to_base64(pdf_path):
        with open(pdf_path, 'rb') as file:
            pdf_data = file.read()
        pdf_base64 = base64.b64encode(pdf_data).decode('utf-8')
        data = {'cv_file': pdf_base64}
        return data
    
    # convert base64 to pdf file
    def base64_to_pdf(base64_string):
        pdf_data = base64.b64decode(base64_string).decode('utf-8')
        with open('resume\last_parsed_cv.pdf', 'wb') as pdf_file:
            pdf_file.write(pdf_data)

    # Converts pdf file to text.
    def pdf_to_text(pdf_path):
        doc = fitz.open(pdf_path)
        text = " "
        for page in doc:
            text = text + str(page.get_text())
        text = text.strip()
        text = " ".join(text.split())
        return text
    
    # Convert word file to text
    def word_to_text(doc_path):
        doc = Document(doc_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return text
    
    # Convert OpenAI Response to Json.
    def convert_string_to_json(json_string):
        try:
            json_object = json.loads(json_string)
            return json_object
        except ValueError as e:
            print("Invalid JSON format:", e)
            return None
    
    # Make Request to OpenAI API with the prompt.
    def make_request_openai(text):
        headers = {
            'Content-Type': 'application/json',
            'Authorization': 'Bearer ' + api_key,
        }
        json_data = {
            'model': 'gpt-3.5-turbo',
            'messages': [
                {
                    'role': 'user',
                    'content': f'{text}\n {q_one_liner}',
                },
            ],
        }
        response = requests.post('https://api.openai.com/v1/chat/completions', headers=headers, json=json_data)
        json_response = json.loads(response.text)
        answers = json_response["choices"][0]["message"]["content"]
        resp = CV_PARSER.convert_string_to_json(answers)
        return resp